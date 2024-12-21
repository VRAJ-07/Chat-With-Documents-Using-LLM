import os
import streamlit as st
from langchain.chains import ConversationalRetrievalChain
from langchain.text_splitter import CharacterTextSplitter
from langchain.vectorstores import FAISS
from langchain_huggingface import HuggingFaceEmbeddings
from langchain.memory import ConversationBufferMemory
from langchain_core.messages import AIMessage, HumanMessage
from langchain.schema import Document
from langchain_huggingface import HuggingFaceEndpoint
import pandas as pd
import json
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
from dotenv import load_dotenv

# Set HuggingFace API Token
load_dotenv()
huggingface_token = os.getenv("HUGGINGFACEHUB_API_TOKEN")

# Function to process different file types
def process_file(file):
    documents = []

    # Process CSV
    if file.name.endswith('.csv'):
        df = pd.read_csv(file)
        for _, row in df.iterrows():
            combined_text = " ".join([str(value) for value in row if not pd.isnull(value)])
            documents.append(Document(page_content=combined_text))

    # Process JSON
    elif file.name.endswith('.json'):
        data = json.load(file)
        if isinstance(data, list):
            for item in data:
                documents.append(Document(page_content=json.dumps(item)))
        else:
            documents.append(Document(page_content=json.dumps(data)))

    # Process PDF
    elif file.name.endswith('.pdf'):
        reader = PdfReader(file)
        text = ""
        for page in reader.pages:
            text += page.extract_text()
        documents.append(Document(page_content=text))

    # Process DOCX
    elif file.name.endswith('.docx'):
        doc = DocxDocument(file)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        documents.append(Document(page_content=text))

    # Split the documents into smaller chunks
    text_splitter = CharacterTextSplitter(separator='\n', chunk_size=1000, chunk_overlap=500)
    split_docs = text_splitter.split_documents(documents)

    # Create Hugging Face embeddings
    huggingface_embeddings = HuggingFaceEmbeddings(model_name='sentence-transformers/all-MiniLM-L6-v2')

    # Create a FAISS vector database
    vectordb = FAISS.from_documents(documents=split_docs, embedding=huggingface_embeddings)
    return vectordb

# Create ChatGroq instance (no dynamic imports)
def create_chat_huggingface():
    return HuggingFaceEndpoint(
         repo_id='mistralai/Mistral-7B-Instruct-v0.2',temperature=0.1
    )

# Streamlit interface
def main():
    # Set page title and icon
    st.set_page_config(page_title="Conversational Chatbot for Multiple Formats", page_icon=":file_folder:")
    st.title("Conversational Chatbot for Multiple File Formats")

    # File uploader
    uploaded_file = st.file_uploader("Upload a file (CSV, JSON, PDF, DOCX):", type=["csv", "json", "pdf", "docx"])

    # Add a submit button
    if st.button("Process File"):
        if uploaded_file:
            # Process file only once
            with st.spinner("Processing the file..."):
                vectordb = process_file(uploaded_file)
            st.success("File processed successfully! Ask your questions below.")

            # Store vector database in session state to avoid re-processing
            st.session_state.vectordb = vectordb

            # Initialize conversation buffer memory for chat history
            memory = ConversationBufferMemory(memory_key="chat_history", return_messages=True)

            # Create a retriever from the FAISS vector database
            retriever = st.session_state.vectordb.as_retriever(search_kwargs={"k": 3})

            # Store retriever and memory in session state separately
            st.session_state.retriever = retriever
            st.session_state.memory = memory

    # Check if the retriever and memory are available before proceeding to the chat
    if "retriever" in st.session_state and "memory" in st.session_state:
        # Initialize or load chat history
        if "chat_history" not in st.session_state:
            st.session_state.chat_history = []

        # Display the chat messages
        for message in st.session_state.chat_history:
            if isinstance(message, AIMessage):
                with st.chat_message("AI"):
                    st.markdown(message.content)
            elif isinstance(message, HumanMessage):
                with st.chat_message("Human"):
                    st.markdown(message.content)

        # User input for questions
        user_query = st.chat_input("Type a message")
        if user_query and user_query.strip() != "":
            # Add human message to the chat history
            st.session_state.chat_history.append(HumanMessage(content=user_query))

            # Display human message
            with st.chat_message("Human"):
                st.markdown(user_query)

            # Get and display the AI response
            with st.chat_message("AI"):
                # Create the ChatGroq model dynamically for each query
                retriever = st.session_state.retriever
                memory = st.session_state.memory

                # Create the ChatGroq instance
                llm = create_chat_huggingface()

                qa_chain = ConversationalRetrievalChain.from_llm(llm=llm, retriever=retriever, memory=memory)
                response = qa_chain({"question": user_query})
                st.markdown(response['answer'])

            # Add AI message to the chat history
            st.session_state.chat_history.append(AIMessage(content=response['answer']))